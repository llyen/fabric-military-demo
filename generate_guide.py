"""
OPERATION IRONSHIELD – Step-by-step Fabric setup guide (DOCX)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT = os.path.dirname(os.path.abspath(__file__))
NAVY  = RGBColor(0x0B, 0x1D, 0x3A)
OLIVE = RGBColor(0x3D, 0x5A, 0x1E)
RED   = RGBColor(0xCC, 0x22, 0x22)
GREY  = RGBColor(0x55, 0x55, 0x55)

def sty(doc):
    s = doc.styles['Normal']; s.font.name = 'Calibri'; s.font.size = Pt(10.5)
    s.font.color.rgb = RGBColor(0x22,0x22,0x22)
    s.paragraph_format.space_after = Pt(6)

def h(doc, txt, lvl=1):
    hd = doc.add_heading(txt, level=lvl)
    for r in hd.runs: r.font.color.rgb = NAVY if lvl<=2 else OLIVE
    return hd

def b(doc, txt, bold_pfx=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_pfx:
        r = p.add_run(bold_pfx); r.bold = True; p.add_run(txt)
    else:
        p.add_run(txt)
    return p

def nb(doc, txt, bold_pfx=None):
    """Numbered list item."""
    p = doc.add_paragraph(style='List Number')
    if bold_pfx:
        r = p.add_run(bold_pfx); r.bold = True; p.add_run(txt)
    else:
        p.add_run(txt)
    return p

def code(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt); r.font.name = 'Consolas'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    return p

def tbl(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,hdr in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = hdr
        for p in c.paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(9)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = v
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    return t

def tip(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run('💡 TIP: '); r.bold = True; r.font.color.rgb = OLIVE
    p.add_run(txt)
    return p

def warn(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run('⚠️ UWAGA: '); r.bold = True; r.font.color.rgb = RED
    p.add_run(txt)
    return p

def nav(doc, txt):
    """UI navigation path."""
    p = doc.add_paragraph()
    r = p.add_run('🖱️  '); r.font.size = Pt(11)
    r2 = p.add_run(txt); r2.bold = True; r2.font.color.rgb = RGBColor(0x00,0x50,0x99)
    r2.font.size = Pt(10)
    return p

# ================================================================

def create_guide():
    doc = Document()
    sty(doc)

    # ── Title ──
    doc.add_paragraph(); doc.add_paragraph()
    t = doc.add_heading('OPERATION IRONSHIELD\nInstrukcja uruchomienia w Microsoft Fabric', level=0)
    for r in t.runs: r.font.color.rgb = NAVY; r.font.size = Pt(24)
    doc.add_paragraph()
    sub = doc.add_paragraph('Krok po kroku: od pustego workspace do działającego demo')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs: r.font.size = Pt(13); r.font.color.rgb = GREY
    doc.add_page_break()

    # ── TOC ──
    h(doc, 'Spis treści', 1)
    toc = [
        'Krok 0: Wymagania wstępne',
        'Krok 1: Utworzenie workspace',
        'Krok 2: Utworzenie Eventhouse i KQL Database',
        'Krok 3: Utworzenie tabel w KQL Database',
        'Krok 4: Utworzenie Eventstream z Custom App source',
        'Krok 5: Upload datasetów do Lakehouse',
        'Krok 6: Uruchomienie generatora danych (notebook)',
        'Krok 7: Konfiguracja Real-Time Dashboard',
        'Krok 8: Konfiguracja Data Agents',
        'Krok 9: Konfiguracja Operations Agents',
        'Krok 10: Test end-to-end i rehearsal',
        'Dodatek A: Troubleshooting',
        'Dodatek B: Ścieżka alternatywna z Azure Event Hub',
    ]
    for t_ in toc: doc.add_paragraph(t_)
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # KROK 0
    # ══════════════════════════════════════════════════
    h(doc, 'Krok 0: Wymagania wstępne', 1)
    doc.add_paragraph(
        'Przed rozpoczęciem konfiguracji upewnij się, że spełnione są następujące warunki:')

    h(doc, '0.1 Licencja i capacity', 2)
    tbl(doc,
        ['Wymaganie', 'Minimum', 'Rekomendowane', 'Uwagi'],
        [
            ['Fabric capacity', 'F2', 'F4 lub wyższa', 'Data Agents wymagają F2+. Trial NIE wspiera Operations Agents.'],
            ['Azure subscription', 'Dowolna', 'Pay-as-you-go', 'Potrzebna tylko jeśli używasz Azure Event Hub (ścieżka alternatywna)'],
            ['Microsoft 365', 'E3/E5 lub Business', 'E5', 'Potrzebny Microsoft Teams do Operations Agent alerts'],
        ]
    )

    h(doc, '0.2 Ustawienia administratora Fabric', 2)
    doc.add_paragraph('W Fabric Admin Portal (admin.fabric.microsoft.com) włącz:')
    nav(doc, 'Admin Portal → Tenant settings → Microsoft Fabric')
    b(doc, ' – WŁĄCZONE', 'Copilot and AI features')
    b(doc, ' – WŁĄCZONE', 'Users can create Fabric items (Eventhouse, Eventstream)')
    b(doc, ' – WŁĄCZONE', 'Real-Time Intelligence')
    b(doc, ' – WŁĄCZONE (dla Data Agents)', 'Data Agent in Microsoft Fabric')
    b(doc, ' – WŁĄCZONE', 'Users can create and share data with Eventstream custom endpoints')
    warn(doc, 'Zmiany tenant settings mogą potrzebować do 15 minut na propagację.')

    h(doc, '0.3 Pliki z tego repozytorium', 2)
    doc.add_paragraph('Upewnij się, że masz folder datasets/ z wygenerowanymi plikami:')
    tbl(doc,
        ['Plik', 'Opis', 'Rozmiar'],
        [
            ['radar_detections.jsonl', 'Wykrycia radarowe (friendly + hostile)', '~3 MB'],
            ['vehicle_status.jsonl', 'Telemetria pojazdów BLUE FORCE', '~9 MB'],
            ['soldier_health.jsonl', 'Dane zdrowotne żołnierzy (wearable)', '~26 MB'],
            ['drone_observations.jsonl', 'Obserwacje z dronów', '~3 MB'],
            ['weather_data.jsonl', 'Dane meteorologiczne', '~1 MB'],
            ['sigint_intercepts.jsonl', 'Przechwycone sygnały SIGINT', '~0.3 MB'],
            ['ammo_logistics.jsonl', 'Logistyka – konwoje i stan amunicji', '~2 MB'],
            ['eventhouse_kql_setup.kql', 'Definicje tabel KQL + mappingi + sample queries', '—'],
            ['fabric_notebook_sender.py', 'Kod notebooka do streamowania danych', '—'],
        ]
    )

    h(doc, '0.4 Narzędzia na komputerze lokalnym (opcjonalnie)', 2)
    b(doc, ' – do lokalnego testowania sendera (pip install azure-eventhub)', 'Python 3.10+')
    b(doc, ' – dowolna (Chrome, Edge) z dostępem do app.fabric.microsoft.com', 'Przeglądarka')
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # KROK 1
    # ══════════════════════════════════════════════════
    h(doc, 'Krok 1: Utworzenie workspace', 1)

    nav(doc, 'app.fabric.microsoft.com → Workspaces (lewe menu) → + New workspace')

    nb(doc, ' „IRONSHIELD-Demo"', 'Nazwa:')
    nb(doc, ' (opcjonalnie) „Operation IRONSHIELD – Fabric RTI battlefield demo"', 'Opis:')
    nb(doc, ' Rozwiń sekcję „Advanced"', 'License mode:')
    b(doc, 'Wybierz Fabric capacity (F2 lub wyższa) – NIE Trial')
    b(doc, 'Potwierdź klikając Apply')

    nb(doc, 'Kliknij Create')

    tip(doc, 'Nazwy workspace w Fabric nie mogą zawierać znaków specjalnych. '
        'Użyj myślników zamiast spacji jeśli wolisz.')
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # KROK 2
    # ══════════════════════════════════════════════════
    h(doc, 'Krok 2: Utworzenie Eventhouse i KQL Database', 1)

    doc.add_paragraph(
        'Eventhouse to kontener dla baz danych KQL. Utworzenie Eventhouse automatycznie '
        'tworzy też domyślną KQL Database o tej samej nazwie.')

    nav(doc, 'Workspace IRONSHIELD-Demo → + New item → wyszukaj „Eventhouse"')

    nb(doc, ' „IRONSHIELD_Eventhouse"', 'Nazwa:')
    nb(doc, 'Kliknij Create')
    nb(doc, 'Po utworzeniu automatycznie otworzy się widok Eventhouse')
    nb(doc, 'W lewym panelu zobaczysz KQL Database „IRONSHIELD_Eventhouse" – to Twoja główna baza')

    tip(doc, 'Eventhouse obsługuje wiele KQL Database, ale dla demo wystarczy jedna domyślna.')
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # KROK 3
    # ══════════════════════════════════════════════════
    h(doc, 'Krok 3: Utworzenie tabel w KQL Database', 1)

    doc.add_paragraph(
        'Otwórz KQL Database i uruchom komendy tworzące tabele. '
        'Użyj pliku eventhouse_kql_setup.kql z folderu datasets/.')

    nav(doc, 'Eventhouse → kliknij KQL Database „IRONSHIELD_Eventhouse" → Explore your data (KQL Queryset)')

    nb(doc, 'W edytorze KQL otwórz / wklej zawartość pliku eventhouse_kql_setup.kql')
    nb(doc, 'Zaznacz PIERWSZY blok .create table (RadarDetections) i kliknij ▶ Run')
    nb(doc, 'Powtórz dla każdej tabeli – łącznie 7 tabel:')

    tbl(doc,
        ['#', 'Tabela', 'Kolumny (kluczowe)'],
        [
            ['1', 'RadarDetections', 'Timestamp, Sector, TrackId, Classification, Lat/Lon, Speed, Confidence'],
            ['2', 'VehicleStatus', 'Timestamp, VehicleId, UnitName, FuelPercent, AmmoPercent, CombatReady'],
            ['3', 'SoldierHealth', 'Timestamp, SoldierId, HeartRate, BodyTemp, BloodOxygen, StressLevel'],
            ['4', 'DroneObservations', 'Timestamp, DroneId, ObservationType, TargetClassification, TargetCount'],
            ['5', 'WeatherData', 'Timestamp, Sector, Visibility_km, CloudCeiling_m, WindSpeed_kmh'],
            ['6', 'SigintIntercepts', 'Timestamp, SignalType, Bearing_deg, Classification, Notes'],
            ['7', 'AmmoLogistics', 'Timestamp, ConvoyId, Status, LoadAmmo_rounds, ETA_min, RouteRisk'],
        ]
    )

    nb(doc, 'Następnie uruchom bloki w kolejności:')
    doc.add_paragraph()
    doc.add_paragraph('Kolejność wykonania bloków z eventhouse_kql_setup.kql:')
    tbl(doc,
        ['Krok', 'Co uruchomić', 'Ile bloków'],
        [
            ['1', '.create table ... (7 tabel + RawEvents)', '8'],
            ['2', '.alter table RawEvents policy retention', '1'],
            ['3', '.create table RawEvents ingestion json mapping', '1'],
            ['4', '.create-or-alter function Extract...() – 7 funkcji', '7'],
            ['5', '.alter table ... policy update – 7 update policies', '7'],
            ['6', '(Opcjonalnie) .create table ... ingestion json mapping per tabela', '3+'],
        ]
    )
    doc.add_paragraph()
    nb(doc, 'Zweryfikuj: w lewym panelu KQL Database powinno widać 8 tabel (7 + RawEvents) i 7 funkcji')

    tip(doc, 'Możesz zaznaczyć cały plik i uruchomić Run all – KQL wykona komendy sekwencyjnie. '
        'Jednak lepiej uruchamiać bloki po kolei, aby łatwiej zidentyfikować ewentualny błąd.')
    warn(doc, 'Jeśli tabela już istnieje, .create zwróci błąd. Użyj .create-or-alter table w takim przypadku. '
         'Update Policies wymagają, aby funkcje Extract*() istniały PRZED ustawieniem policy.')
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # KROK 4
    # ══════════════════════════════════════════════════
    h(doc, 'Krok 4: Utworzenie Eventstream z Custom App source', 1)

    doc.add_paragraph(
        'Eventstream odbiera zdarzenia i przekazuje je do Eventhouse. '
        'Użyjemy Custom App source – Fabric utworzy endpoint kompatybilny z Event Hub SDK, '
        'bez potrzeby osobnego Azure Event Hub.')

    h(doc, '4.1 Utwórz Eventstream', 2)
    nav(doc, 'Workspace → + New item → wyszukaj „Eventstream" → Create')
    nb(doc, ' „IRONSHIELD-Stream"', 'Nazwa:')
    nb(doc, 'Kliknij Create – otworzy się canvas Eventstream')

    h(doc, '4.2 Dodaj Custom App source', 2)
    nav(doc, 'Canvas → + Add source (lewy panel) → Custom App')
    nb(doc, ' „ironshield-sensor-input"', 'Source name:')
    nb(doc, 'Kliknij Add')
    nb(doc, 'Po dodaniu, kliknij na source „ironshield-sensor-input" na canvasie')
    nb(doc, 'W panelu po prawej zobaczysz: Connection string i Event Hub name')
    nb(doc, '', 'SKOPIUJ I ZAPISZ connection string – będzie potrzebny w kroku 6!')

    warn(doc, 'Connection string jest specyficzny dla tego Eventstreamʼa. '
         'Nie mylić z Azure Event Hub connection string. Format: '
         'Endpoint=sb://....servicebus.windows.net/;SharedAccessKeyName=...;SharedAccessKey=...;EntityPath=...')

    h(doc, '4.3 Dodaj destination (KQL Database – tabela RawEvents)', 2)
    doc.add_paragraph(
        'W naszym podejściu WSZYSTKIE zdarzenia (z 7 strumieni) trafiają do jednej tabeli RawEvents. '
        'Routing do tabel docelowych realizują KQL Update Policies (ustawione w Kroku 3).')

    h(doc, 'Podejście A: Jeden Eventstream + KQL Update Policies (ZALECANE)', 3)
    doc.add_paragraph(
        'Jest to najczystsze podejście. Eventstream nie wymaga żadnych transformacji – wystarczy '
        'jeden source i jedna destination. Routing odbywa się po stronie Eventhouse dzięki '
        'Update Policies, które automatycznie parsują dane z RawEvents i wstawiają do 7 tabel.')

    nav(doc, 'Canvas → + Add destination → Eventhouse')
    nb(doc, ' „dest-raw-events"', 'Destination name:')

    warn(doc, 'KLUCZOWY KROK: Wybierz tryb ingestion „Direct ingestion" (NIE „Event processing '
         'before ingestion"). Tryb „Event processing" próbuje dopasować pola JSON do kolumn tabeli '
         'i zwróci błąd „request is invalid", bo RawEvents ma tylko jedną kolumnę dynamic.')

    doc.add_paragraph()
    doc.add_paragraph('Konfiguracja destination:')
    nb(doc, ' Direct ingestion', 'Ingestion mode:')
    nb(doc, ' IRONSHIELD-Demo (Twój workspace)', 'Workspace:')
    nb(doc, ' IRONSHIELD_Eventhouse', 'Eventhouse:')
    nb(doc, ' IRONSHIELD_Eventhouse (domyślna KQL Database)', 'KQL Database:')
    nb(doc, ' RawEvents', 'Table:')
    nb(doc, ' JSON', 'Input data format:')
    nb(doc, ' RawEvents_mapping (utworzony w Kroku 3)', 'Ingestion mapping:')
    nb(doc, 'Kliknij Add')

    doc.add_paragraph()
    doc.add_paragraph('Schemat przepływu danych:')
    code(doc,
        'Notebook (JSONL) → Event Hub SDK → Custom App Source\n'
        '    → Eventstream → RawEvents (Eventhouse)\n'
        '        → Update Policy → RadarDetections\n'
        '        → Update Policy → VehicleStatus\n'
        '        → Update Policy → SoldierHealth\n'
        '        → Update Policy → DroneObservations\n'
        '        → Update Policy → WeatherData\n'
        '        → Update Policy → SigintIntercepts\n'
        '        → Update Policy → AmmoLogistics'
    )

    doc.add_paragraph()
    tip(doc, 'Update Policies działają automatycznie – za każdym razem gdy dane trafią do RawEvents, '
        'KQL Engine uruchamia funkcje Extract*() i wstawia wyniki do tabel docelowych. '
        'Opóźnienie to zazwyczaj < 1 sekunda.')

    doc.add_paragraph()
    warn(doc, 'WAŻNE: tabela RawEvents ma ustawiony retention na 1 dzień. '
         'Surowe dane są usuwane automatycznie – dane docelowe pozostają w 7 tabelach. '
         'W razie potrzeby zmień retention: .alter table RawEvents policy retention ...')

    h(doc, 'Podejście B: 7 osobnych Eventstreamów (fallback)', 3)
    doc.add_paragraph(
        'Jeśli z jakiegokolwiek powodu Update Policies nie działają lub chcesz uprościć '
        'debugging, możesz utworzyć 7 osobnych Eventstreamów – po jednym per strumień. '
        'Każdy Eventstream ma własny Custom App source i destination do odpowiedniej tabeli. '
        'Notebook sender musi wtedy mieć 7 connection stringów.')

    tbl(doc,
        ['Eventstream', 'Custom App Source', 'Destination table'],
        [
            ['IRONSHIELD-Radar', 'radar-input', 'RadarDetections'],
            ['IRONSHIELD-Vehicle', 'vehicle-input', 'VehicleStatus'],
            ['IRONSHIELD-Soldier', 'soldier-input', 'SoldierHealth'],
            ['IRONSHIELD-Drone', 'drone-input', 'DroneObservations'],
            ['IRONSHIELD-Weather', 'weather-input', 'WeatherData'],
            ['IRONSHIELD-SIGINT', 'sigint-input', 'SigintIntercepts'],
            ['IRONSHIELD-Logistics', 'logistics-input', 'AmmoLogistics'],
        ]
    )

    tip(doc, 'Podejście B wymaga 7 connection stringów i 7 Eventstreamów, ale nie wymaga '
        'Update Policies ani tabeli RawEvents. Przy tym podejściu pomiń RawEvents w Kroku 3.')

    h(doc, '4.4 Publish Eventstream', 2)
    nav(doc, 'Canvas → Publish (prawy górny róg)')
    doc.add_paragraph('Po opublikowaniu Eventstream jest aktywny i czeka na dane.')
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # KROK 5
    # ══════════════════════════════════════════════════
    h(doc, 'Krok 5: Upload datasetów do Lakehouse', 1)

    doc.add_paragraph(
        'Pliki JSONL muszą być dostępne z poziomu Fabric notebooka. '
        'Najprościej uploadować je do Lakehouse w workspace.')

    h(doc, '5.1 Utwórz Lakehouse', 2)
    nav(doc, 'Workspace → + New item → Lakehouse')
    nb(doc, ' „IRONSHIELD_Lakehouse"', 'Nazwa:')
    nb(doc, 'Kliknij Create')

    h(doc, '5.2 Upload plików', 2)
    nav(doc, 'Lakehouse → Files (lewy panel) → Upload → Upload folder')
    nb(doc, 'Wybierz folder datasets/ z plikiami JSONL')
    nb(doc, 'Lub: Upload → Upload files → zaznacz 7 plików .jsonl')
    nb(doc, 'Poczekaj na zakończenie uploadu (pliki łącznie ~44 MB)')
    nb(doc, 'Zweryfikuj: w sekcji Files/ powinno widać 7 plików .jsonl')

    tip(doc, 'Ścieżka do plików w notebooku: '
        'abfss://IRONSHIELD-Demo@onelake.dfs.fabric.microsoft.com/IRONSHIELD_Lakehouse.Lakehouse/Files/')
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # KROK 6
    # ══════════════════════════════════════════════════
    h(doc, 'Krok 6: Uruchomienie generatora danych (notebook)', 1)

    doc.add_paragraph(
        'Notebook czyta pliki JSONL z Lakehouse i wysyła je do Eventstream Custom App '
        'w odpowiednich odstępach czasowych, symulując dane real-time z sensorów.')

    h(doc, '6.1 Utwórz Notebook', 2)
    nav(doc, 'Workspace → + New item → Notebook')
    nb(doc, ' „IRONSHIELD_DataSender"', 'Nazwa:')
    nb(doc, 'Zmień język na PySpark (Python) jeśli trzeba')

    h(doc, '6.2 Zainstaluj SDK', 2)
    doc.add_paragraph('W pierwszej komórce notebooka:')
    code(doc, '%pip install azure-eventhub --quiet')
    doc.add_paragraph('Uruchom komórkę i poczekaj na instalację.')

    h(doc, '6.3 Konfiguracja', 2)
    doc.add_paragraph('W drugiej komórce wklej konfigurację – zaktualizuj wartości:')
    code(doc,
        '# ═══ KONFIGURACJA ═══\n'
        '# Connection string z Eventstream Custom App source (Krok 4.2)\n'
        'EVENT_HUB_CONNECTION_STRING = "Endpoint=sb://....;SharedAccessKeyName=...;SharedAccessKey=...;EntityPath=..."\n'
        '\n'
        '# Ścieżka do plików JSONL w Lakehouse\n'
        'LAKEHOUSE_PATH = "abfss://IRONSHIELD-Demo@onelake.dfs.fabric.microsoft.com/IRONSHIELD_Lakehouse.Lakehouse/Files/"\n'
        '\n'
        '# Prędkość symulacji: 1.0 = real-time (20 min), 0.5 = 2x szybciej (10 min)\n'
        'TIME_SCALE = 0.5\n'
        '\n'
        '# Rozmiar batcha do wysyłki\n'
        'BATCH_SIZE = 100'
    )

    warn(doc, 'Connection string MUSI być z Custom App source Eventstream (krok 4.2), '
         'NIE z Azure Event Hub.')

    h(doc, '6.4 Kod sendera', 2)
    doc.add_paragraph(
        'W trzeciej komórce wklej poniższy kod (lub użyj pliku fabric_notebook_sender.py '
        'z folderu datasets/):')
    code(doc,
        'import json, time\n'
        'from datetime import datetime\n'
        'from azure.eventhub import EventHubProducerClient, EventData\n'
        '\n'
        '# Load all events from JSONL files\n'
        'STREAMS = [\n'
        '    "radar_detections", "vehicle_status", "soldier_health",\n'
        '    "drone_observations", "weather_data", "sigint_intercepts", "ammo_logistics"\n'
        ']\n'
        '\n'
        'all_events = []\n'
        'for stream_name in STREAMS:\n'
        '    path = f"{LAKEHOUSE_PATH}{stream_name}.jsonl"\n'
        '    df = spark.read.text(path)\n'
        '    for row in df.collect():\n'
        '        event = json.loads(row.value)\n'
        '        event["_stream"] = stream_name\n'
        '        all_events.append(event)\n'
        '    print(f"  Loaded {stream_name}: {df.count()} events")\n'
        '\n'
        'all_events.sort(key=lambda e: e.get("Timestamp", ""))\n'
        'print(f"\\nTotal events: {len(all_events):,}")\n'
        '\n'
        '# Send to Eventstream\n'
        'producer = EventHubProducerClient.from_connection_string(\n'
        '    conn_str=EVENT_HUB_CONNECTION_STRING\n'
        ')\n'
        '\n'
        'prev_ts = None\n'
        'sent = 0\n'
        'batch = producer.create_batch()\n'
        'batch_count = 0\n'
        '\n'
        'print("\\n🚀 Streaming started...")\n'
        'for event in all_events:\n'
        '    ts = event.get("Timestamp", "")\n'
        '    stream = event.pop("_stream", "unknown")\n'
        '\n'
        '    # Simulate real-time delay\n'
        '    if prev_ts and ts:\n'
        '        try:\n'
        '            delta = (datetime.fromisoformat(ts) - datetime.fromisoformat(prev_ts)).total_seconds()\n'
        '            time.sleep(min(delta * TIME_SCALE, 3.0))\n'
        '        except: pass\n'
        '    prev_ts = ts\n'
        '\n'
        '    ed = EventData(json.dumps(event))\n'
        '    ed.properties = {"stream": stream}\n'
        '    try:\n'
        '        batch.add(ed)\n'
        '        batch_count += 1\n'
        '    except ValueError:\n'
        '        producer.send_batch(batch)\n'
        '        sent += batch_count\n'
        '        batch = producer.create_batch()\n'
        '        batch.add(ed)\n'
        '        batch_count = 1\n'
        '\n'
        '    if batch_count >= BATCH_SIZE:\n'
        '        producer.send_batch(batch)\n'
        '        sent += batch_count\n'
        '        batch = producer.create_batch()\n'
        '        batch_count = 0\n'
        '        if sent % 5000 == 0:\n'
        '            print(f"  Sent {sent:,} / {len(all_events):,}")\n'
        '\n'
        'if batch_count > 0:\n'
        '    producer.send_batch(batch)\n'
        '    sent += batch_count\n'
        '\n'
        'producer.close()\n'
        'print(f"\\n✅ Done! Sent {sent:,} events.")'
    )

    h(doc, '6.5 Uruchomienie', 2)
    nb(doc, 'Uruchom komórki po kolei: install → config → sender')
    nb(doc, 'Sender będzie wyświetlał postęp: "Sent 5,000 / 138,100" itd.')
    nb(doc, 'Przy TIME_SCALE=0.5 symulacja potrwa ~10 minut')
    nb(doc, 'W trakcie przejdź do dashboardu (krok 7) i obserwuj napływ danych')

    tip(doc, 'Możesz też uruchomić sender lokalnie z terminala (Python + pip install azure-eventhub). '
        'W takim przypadku zamień spark.read.text() na zwykły open() dla plików .jsonl.')
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # KROK 7
    # ══════════════════════════════════════════════════
    h(doc, 'Krok 7: Konfiguracja Real-Time Dashboard', 1)

    h(doc, '7.1 Utwórz Dashboard', 2)
    nav(doc, 'Workspace → + New item → Real-Time Dashboard')
    nb(doc, ' „IRONSHIELD – Battlefield COP"', 'Nazwa:')
    nb(doc, 'Kliknij Create')

    h(doc, '7.2 Podłącz data source', 2)
    nav(doc, 'Dashboard → + New data source → OneLake data hub → wybierz IRONSHIELD_Eventhouse')
    nb(doc, 'Kliknij Connect')
    nb(doc, 'Teraz masz dostęp do wszystkich 7 tabel KQL')

    h(doc, '7.3 Dodaj tile: Mapa operacyjna (Scatter Map)', 2)
    nav(doc, 'Dashboard → + Add tile')
    nb(doc, 'Wklej query:')
    code(doc,
        'RadarDetections\n'
        '| where Timestamp > ago(5m)\n'
        '| project Latitude, Longitude, Classification, ObjectType, Confidence\n'
        '| extend Color = case(\n'
        '    Classification == "hostile", "red",\n'
        '    Classification == "friendly", "blue",\n'
        '    "gray")')
    nb(doc, 'Visualization type: Map')
    nb(doc, 'Latitude: Latitude, Longitude: Longitude, Color by: Color')
    nb(doc, 'Auto-refresh: 5 seconds')
    nb(doc, 'Save tile')

    h(doc, '7.4 Dodaj tile: Status sił BLUE FORCE (Stat/KPI)', 2)
    code(doc,
        'VehicleStatus\n'
        '| where Timestamp > ago(2m)\n'
        '| summarize arg_max(Timestamp, *) by VehicleId\n'
        '| summarize \n'
        '    ReadyCount = countif(CombatReady),\n'
        '    TotalCount = count(),\n'
        '    AvgAmmo = round(avg(AmmoPercent),0),\n'
        '    AvgFuel = round(avg(FuelPercent),0)\n'
        '  by UnitName')
    nb(doc, 'Visualization: Table lub Multi-stat')
    nb(doc, 'Auto-refresh: 10 seconds')

    h(doc, '7.5 Dodaj tile: Wykrycia hostile – oś czasu', 2)
    code(doc,
        'RadarDetections\n'
        '| where Classification == "hostile"\n'
        '| summarize ContactCount = dcount(TrackId) by bin(Timestamp, 30s), Sector\n'
        '| render timechart')
    nb(doc, 'Visualization: Time chart (auto-detected)')
    nb(doc, 'Auto-refresh: 5 seconds')

    h(doc, '7.6 Dodaj tile: Soldier health alerts', 2)
    code(doc,
        'SoldierHealth\n'
        '| where Timestamp > ago(3m)\n'
        '| summarize arg_max(Timestamp, *) by SoldierId\n'
        '| where HeartRate > 160 or BodyTemp > 38.5 or BloodOxygen < 93\n'
        '| project SoldierId, UnitName, Sector, HeartRate, BodyTemp, BloodOxygen, StressLevel\n'
        '| order by StressLevel asc')
    nb(doc, 'Visualization: Table')

    h(doc, '7.7 Dodaj tile: Logistyka – stan amunicji', 2)
    code(doc,
        'AmmoLogistics\n'
        '| where CargoType == "status_report" and Timestamp > ago(2m)\n'
        '| summarize arg_max(Timestamp, *) by DestinationUnit\n'
        '| project DestinationUnit, UnitAmmoPercent, UnitFuelPercent, ResupplyNeeded\n'
        '| order by UnitAmmoPercent asc')
    nb(doc, 'Visualization: Bar chart lub Table')

    h(doc, '7.8 Dodaj tile: Pogoda operacyjna', 2)
    code(doc,
        'WeatherData\n'
        '| where Timestamp > ago(5m)\n'
        '| summarize arg_max(Timestamp, *) by Sector\n'
        '| project Sector, Visibility_km, CloudCeiling_m, WindSpeed_kmh, Precipitation')

    h(doc, '7.9 Dodaj tile: SIGINT activity', 2)
    code(doc,
        'SigintIntercepts\n'
        '| where Timestamp > ago(10m)\n'
        '| summarize Count = count() by bin(Timestamp, 1m), Classification\n'
        '| render timechart')

    h(doc, '7.10 Konfiguracja globalna dashboardu', 2)
    nb(doc, 'Ustaw auto-refresh dashboardu: 5 sekund (minimum)')
    nav(doc, 'Dashboard → ⚙️ Settings → Auto refresh → Every 5 seconds')
    nb(doc, 'Ustaw time range: Last 20 minutes')
    nb(doc, 'Opcjonalnie: dodaj parametr „Sector" do filtrowania per sektor')

    tip(doc, 'Rozmieść tile w 2-3 kolumnach. Mapa operacyjna powinna być największym elementem '
        '(lewy górny róg, ~50% szerokości). Alertowe tile na prawo.')
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # KROK 8
    # ══════════════════════════════════════════════════
    h(doc, 'Krok 8: Konfiguracja Data Agents', 1)

    doc.add_paragraph(
        'Data Agents umożliwiają dowódcy (użytkownikowi) zadawanie pytań w języku naturalnym '
        'i otrzymywanie odpowiedzi z danych KQL Database.')

    h(doc, '8.1 Utwórz Data Agent – „IRONSHIELD Intel"', 2)
    nav(doc, 'Workspace → + New item → wyszukaj „Data Agent" → Create')
    nb(doc, ' „IRONSHIELD Intel"', 'Nazwa:')
    nb(doc, 'Kliknij Create – otworzy się panel konfiguracji agenta')

    h(doc, '8.2 Dodaj źródła danych', 2)
    nav(doc, 'Agent config → + Add data source → wybierz „KQL Database"')
    nb(doc, 'Wybierz IRONSHIELD_Eventhouse → IRONSHIELD_Eventhouse (KQL DB)')
    nb(doc, 'Zaznacz tabele, do których agent ma mieć dostęp:')
    b(doc, 'RadarDetections ✓')
    b(doc, 'DroneObservations ✓')
    b(doc, 'SigintIntercepts ✓')
    b(doc, 'WeatherData ✓')
    nb(doc, 'Kliknij Add / Confirm')

    h(doc, '8.3 Ustaw instrukcje agenta', 2)
    doc.add_paragraph('W sekcji „Instructions" (system prompt) wpisz:')
    code(doc,
        'Jesteś agentem wywiadu wojskowego IRONSHIELD Intel.\n'
        'Odpowiadasz na pytania dotyczące sytuacji na polu walki.\n'
        'Twoje źródła danych to: wykrycia radarowe, obserwacje z dronów,\n'
        'przechwycone sygnały SIGINT i dane pogodowe.\n'
        '\n'
        'Zawsze podawaj:\n'
        '- Liczbę wykrytych kontaktów\n'
        '- Klasyfikację (hostile/friendly/unknown)\n'
        '- Kierunek i prędkość ruchu\n'
        '- Szacowany czas dotarcia (ETA) do pozycji BLUE FORCE\n'
        '- Poziom pewności (confidence) danych\n'
        '\n'
        'Odpowiadaj zwięźle, w stylu raportów wojskowych.\n'
        'Koreluj dane z wielu tabel gdy to możliwe.')

    h(doc, '8.4 Testuj agenta', 2)
    doc.add_paragraph('W panelu czatu po prawej stronie wpisz pytania testowe:')
    b(doc, '"Ile kontaktów hostile wykryto w sektorze Bravo w ostatnich 10 minutach?"')
    b(doc, '"Jaki jest kierunek i prędkość ruchu wrogiej kolumny?"')
    b(doc, '"Czy warunki pogodowe pozwalają na operacje dronów w sektorze Bravo?"')
    b(doc, '"Podsumuj aktywność SIGINT z ostatnich 5 minut."')

    h(doc, '8.5 Utwórz pozostałe agenty', 2)
    doc.add_paragraph('Powtórz kroki 8.1–8.4 dla dwóch dodatkowych agentów:')
    tbl(doc,
        ['Agent', 'Tabele źródłowe', 'Rola w instrukcjach'],
        [
            ['IRONSHIELD Supply', 'VehicleStatus, AmmoLogistics, WeatherData',
             'Agent logistyczny: stan amunicji/paliwa, ETA konwojów, gotowość bojowa'],
            ['IRONSHIELD Medic', 'SoldierHealth, VehicleStatus',
             'Agent medyczny: monitoring zdrowia żołnierzy, alerty MEDEVAC, lokalizacja rannych'],
        ]
    )

    tip(doc, 'Data Agents potrzebują danych w tabelach, żeby mogły odpowiadać. '
        'Uruchom sender (krok 6) ZANIM testujesz agentów.')
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # KROK 9
    # ══════════════════════════════════════════════════
    h(doc, 'Krok 9: Konfiguracja Operations Agents', 1)

    doc.add_paragraph(
        'Operations Agents monitorują dane w Eventhouse i automatycznie rekomendują akcje '
        'w cyklu observe → analyze → decide → act.')

    h(doc, '9.1 Utwórz Operations Agent – „Threat Response"', 2)
    nav(doc, 'Workspace → + New item → wyszukaj „Operations Agent" (w kategorii Real-Time Intelligence) → Create')
    nb(doc, ' „IRONSHIELD Threat Response"', 'Nazwa:')
    nb(doc, 'Kliknij Create – otworzy się konfiguracja agenta')

    h(doc, '9.2 Skonfiguruj Business Goal', 2)
    doc.add_paragraph('W sekcji „Business Goal" opisz cel agenta:')
    code(doc,
        'Monitoruj wykrycia radarowe w poszukiwaniu koncentracji sił wroga.\n'
        'Jeśli w dowolnym sektorze wykryjesz 10+ kontaktów hostile zbliżających się\n'
        'do pozycji BLUE FORCE (dystans < 15 km), zarekomenduj ogień zaporowy\n'
        'artylerii z baterii Krab wraz ze współrzędnymi celu.')

    h(doc, '9.3 Podłącz Knowledge Source', 2)
    nav(doc, 'Agent config → Knowledge source → wybierz IRONSHIELD_Eventhouse (KQL DB)')
    nb(doc, 'Zaznacz tabele: RadarDetections, DroneObservations, VehicleStatus')

    h(doc, '9.4 Skonfiguruj Recommended Actions', 2)
    doc.add_paragraph('W sekcji „Actions" dodaj:')
    nb(doc, ' „Alarmuj baterię artylerii" – typ: Teams notification', 'Akcja 1:')
    nb(doc, ' „Wyślij dron BDA do weryfikacji" – typ: Teams notification', 'Akcja 2:')

    tip(doc, 'Na tym etapie Operations Agents jest w preview. Jeśli w Twoim tenancie nie jest jeszcze '
        'dostępny jako osobny item, możesz symulować ten krok używając Fabric Activator (Data Activator) '
        'który oferuje podobny cykl: trigger → condition → action.')

    h(doc, '9.5 Activator jako alternatywa (jeśli Operations Agent niedostępny)', 2)
    nav(doc, 'Workspace → + New item → Reflex (Activator)')
    nb(doc, 'Podłącz do Eventstream lub KQL Database')
    nb(doc, 'Zdefiniuj warunek:')
    code(doc,
        'RadarDetections\n'
        '| where Classification == "hostile"\n'
        '| summarize ContactCount = dcount(TrackId) by Sector\n'
        '| where ContactCount >= 10')
    nb(doc, 'Akcja: Send Teams message / Send email / Start Power Automate flow')

    h(doc, '9.6 Utwórz pozostałe Operations Agents / Activators', 2)
    tbl(doc,
        ['Agent', 'Cel', 'Warunek', 'Akcja'],
        [
            ['Logistics Optimizer', 'Utrzymanie zdolności bojowej',
             'AmmoLogistics: UnitAmmoPercent < 30',
             'Teams: alert o niskim stanie + rekomendacja konwoju'],
            ['MEDEVAC Coordinator', 'Ochrona życia żołnierzy',
             'SoldierHealth: HeartRate > 200 OR BloodOxygen < 90',
             'Teams: natychmiastowy alert MEDEVAC + lokalizacja'],
            ['EW Alert', 'Ochrona łączności',
             'SigintIntercepts: SignalType == "jamming_signal"',
             'Automatyczna notyfikacja (pre-approved)'],
        ]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # KROK 10
    # ══════════════════════════════════════════════════
    h(doc, 'Krok 10: Test end-to-end i rehearsal', 1)

    h(doc, '10.1 Checklist przed uruchomieniem demo', 2)
    b(doc, '☐ Eventhouse z 8 tabelami (7 docelowych + RawEvents) + 7 Update Policies')
    b(doc, '☐ Eventstream opublikowany, Custom App source → RawEvents destination')
    b(doc, '☐ Lakehouse z 7 plikami JSONL')
    b(doc, '☐ Notebook z poprawnym connection string')
    b(doc, '☐ Real-Time Dashboard z 7 tile\'ami, auto-refresh 5s')
    b(doc, '☐ 3 Data Agents skonfigurowane i przetestowane')
    b(doc, '☐ Operations Agents / Activators skonfigurowane')
    b(doc, '☐ Teams kanał „IRONSHIELD-OPS" z powiadomieniami')

    h(doc, '10.2 Scenariusz testu', 2)
    nb(doc, 'Otwórz Real-Time Dashboard na jednym ekranie')
    nb(doc, 'Otwórz notebook sender na drugim ekranie (lub w tle)')
    nb(doc, 'Uruchom sender (TIME_SCALE=0.5 → ~10 min)')
    nb(doc, 'Obserwuj dashboard – dane powinny pojawić się w ciągu 5-10 sekund')
    nb(doc, 'Po ~2 min: otwórz Data Agent IRONSHIELD Intel → zadaj pytanie testowe')
    nb(doc, 'Po ~4 min: na dashboardzie powinien pojawić się spike hostile contacts')
    nb(doc, 'Po ~5 min: Operations Agent / Activator powinien wysłać alert w Teams')
    nb(doc, 'Po ~8 min: sprawdź logistykę – 3. kompania poniżej 30% amunicji')

    h(doc, '10.3 Reset danych między próbami', 2)
    doc.add_paragraph('Aby wyczyścić dane i uruchomić symulację od nowa:')
    code(doc,
        '// Uruchom w KQL Database:\n'
        '.clear table RawEvents data\n'
        '.clear table RadarDetections data\n'
        '.clear table VehicleStatus data\n'
        '.clear table SoldierHealth data\n'
        '.clear table DroneObservations data\n'
        '.clear table WeatherData data\n'
        '.clear table SigintIntercepts data\n'
        '.clear table AmmoLogistics data')

    tip(doc, 'Podczas demo na żywo: miej tę komendę gotową w osobnym KQL Queryset oknie, '
        'żebyś mógł szybko zresetować dane między sesjami.')
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # DODATEK A
    # ══════════════════════════════════════════════════
    h(doc, 'Dodatek A: Troubleshooting', 1)

    tbl(doc,
        ['Problem', 'Przyczyna', 'Rozwiązanie'],
        [
            ['Dane nie pojawiają się w KQL',
             'Connection string niepoprawny',
             'Sprawdź czy skopiowałeś CAŁY connection string z Custom App source (z EntityPath)'],
            ['Dane w KQL ale dashboard pusty',
             'Time range mismatch',
             'Dashboard time range musi obejmować czas eventów. Ustaw „Last 30 min" lub „Last 1 hour"'],
            ['Data Agent nie odpowiada',
             'Brak danych w tabelach',
             'Uruchom sender ZANIM testujesz agenta. Sprawdź count() w KQL.'],
            ['Data Agent: „I cannot access..."',
             'Agent nie ma dostępu do tabel',
             'Wróć do konfiguracji agenta i zweryfikuj mapowanie tabel'],
            ['Eventstream: „Publishing failed"',
             'Konfiguracja destination niekompletna',
             'Sprawdź czy tabela destination istnieje i mapping jest poprawny'],
            ['Notebook: „Connection refused"',
             'Firewall / network',
             'Sprawdź czy Fabric compute ma dostęp do endpoint Eventstream. Zazwyczaj działa out-of-the-box.'],
            ['Operations Agent niedostępny',
             'Feature w preview / nie włączony',
             'Sprawdź tenant settings. Użyj Activator (Reflex) jako alternatywę.'],
            ['Dane trafiają do RawEvents ale nie do tabel docelowych',
             'Update Policy nie działa lub funkcja Extract ma błąd',
             'Sprawdź: .show table RadarDetections policy update — IsEnabled musi być true. '
             'Uruchom Extract*() ręcznie w KQL aby zweryfikować wynik.'],
            ['Dane w tabelach są zduplikowane',
             'Update Policy + bezpośredni ingestion jednocześnie',
             'Użyj ALBO Update Policy (podejście A) ALBO direct ingestion (podejście B), nie oba.'],
        ]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # DODATEK B
    # ══════════════════════════════════════════════════
    h(doc, 'Dodatek B: Ścieżka alternatywna z Azure Event Hub', 1)

    doc.add_paragraph(
        'Jeśli wolisz użyć osobnego Azure Event Hub (np. masz już istniejącą infrastrukturę '
        'lub chcesz wysyłać dane spoza Fabric), wykonaj te dodatkowe kroki:')

    h(doc, 'B.1 Utwórz Event Hub w Azure Portal', 2)
    nav(doc, 'portal.azure.com → Create resource → Event Hubs → Create')
    nb(doc, 'Namespace: „ironshield-eventhub-ns" (globalnie unikalna nazwa)')
    nb(doc, 'Pricing tier: Basic (wystarczający dla demo)')
    nb(doc, 'Po utworzeniu namespace: + Event Hub → Nazwa: „ironshield-events"')
    nb(doc, 'Shared Access Policies → + Add → Nazwa: „sender" → Manage ✓')
    nb(doc, 'Skopiuj Primary Connection String')

    h(doc, 'B.2 Zmień źródło w Eventstream', 2)
    doc.add_paragraph('Zamiast Custom App source, użyj Azure Event Hub source:')
    nav(doc, 'Eventstream canvas → + Add source → Azure Event Hubs')
    nb(doc, 'Podaj connection string z Azure Portal')
    nb(doc, 'Consumer group: $Default lub utwórz nowy')
    nb(doc, 'Reszta konfiguracji identyczna jak w kroku 4')

    h(doc, 'B.3 Zaktualizuj connection string w notebooku', 2)
    doc.add_paragraph('W notebooku (krok 6) zamień connection string na ten z Azure Event Hub. '
        'Dodaj też EVENT_HUB_NAME:')
    code(doc,
        'EVENT_HUB_CONNECTION_STRING = "<connection string z Azure Portal>"\n'
        'EVENT_HUB_NAME = "ironshield-events"  # potrzebny przy Azure Event Hub\n'
        '\n'
        'producer = EventHubProducerClient.from_connection_string(\n'
        '    conn_str=EVENT_HUB_CONNECTION_STRING,\n'
        '    eventhub_name=EVENT_HUB_NAME  # dodaj ten parametr\n'
        ')')

    tip(doc, 'Custom App source (krok 4) jest prostszy i nie wymaga Azure subscription. '
        'Azure Event Hub jest lepszy gdy wysyłasz dane z zewnętrznych systemów lub z wielu lokalizacji.')

    # ── Save ──
    fname = 'IRONSHIELD_Setup_Guide_v2.docx'
    path = os.path.join(OUT, fname)
    doc.save(path)
    print(f"Created: {path}")

if __name__ == '__main__':
    create_guide()
